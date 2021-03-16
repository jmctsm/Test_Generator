#!python

"""
Generates a matching file 
will use a .txt file in the current directory
expects the file to be key\n\tdef
"""

import random
import string
from typing import KeysView
import os
import re
from docx import Document


def random_letter():
    if len(key_words) > 26:
        return f"{random.choice(string.ascii_uppercase)}{random.choice(string.ascii_uppercase)}"
    else:
        return random.choice(string.ascii_uppercase)


def readfile():
    """
    Read in all .txt files in the main directory to create the study list
    """
    entries = os.listdir(os.path.dirname(os.path.realpath(__file__)))
    print(entries)
    export_files = []
    doc_files = 0
    for file_name in entries:
        if file_name.endswith(".txt"):
            export_files.append(file_name)
        if file_name.endswith(".docx"):
            doc_files += 1
    if len(export_files) > 1:
        raise ValueError(f"Too many .txt files in the directory.")
    for file in export_files:
        print(f"Using the file {file} to make the quiz")
        with open(file, "r") as export_file:
            data = export_file.readlines()
        keys = []
        defs = []
        for line in data:
            if line == "\t\n" or line == "\n":
                continue
            if line.startswith("\t"):
                line = re.sub(r"[\t\n]", "", line)
                defs.append(line)
            else:
                line = re.sub(r"[\t\n]", "", line)
                keys.append(line)
        doc_file_name = file[:-4]
        save_file = f"{doc_file_name}_{doc_files + 1}.docx"
    return (keys, defs, save_file)


if __name__ == "__main__":
    print("Starting to get the info from the files")
    key_words, definitions, doc_file = readfile()
    test_dict = {}
    answer_list = {}
    counter = 0
    used_numbers = []
    print("Working on creating the quiz")
    while True:
        if len(test_dict) > len(key_words):
            break
        letter = random_letter()
        if letter in test_dict.keys():
            continue
        random_number = random.randint(0, len(key_words) - 1)
        if random_number in used_numbers:
            continue
        used_numbers.append(random_number)
        test_dict[letter] = definitions[random_number]
        answer_list[random_number] = letter
        if len(test_dict) == len(key_words):
            break
        counter += 1
    print("Quiz made.  Now to make the output")
    print("\n\nQUIZ")
    counter = 0
    for key in key_words:
        print(f"{counter + 1}.  _____{key_words[counter]}")
        counter += 1
    print("\n\n\n")
    for key in test_dict.keys():
        print(f"{key}.  {test_dict[key]}")
    print("\n\n\nANSWERS")
    for x in range(len(answer_list)):
        print(f"{x + 1} {key_words[x]} : {answer_list[x]}")

    # Create a word doc
    print("\n\nCreating the Word doc.")
    document = Document()
    document.add_paragraph("QUIZ")
    counter = 0
    for key in key_words:
        document.add_paragraph(f"{counter + 1}.  _____{key_words[counter]}")
        counter += 1
    document.add_paragraph("\n")
    document.add_paragraph("\n")
    document.add_paragraph("\n")
    for key in test_dict.keys():
        document.add_paragraph(f"{key}.  {test_dict[key]}")
    document.add_paragraph("\n")
    document.add_paragraph("\n")
    document.add_paragraph("\n")
    document.add_page_break()
    document.add_paragraph("ANSWERS")
    for x in range(len(answer_list)):
        document.add_paragraph(f"{x + 1} {key_words[x]} : {answer_list[x]}")
    document.save(doc_file)
    print(f"Quiz file created at {doc_file}")
